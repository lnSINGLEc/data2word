##好评占比
##物流方式占比
##总评论数按时间分布
##各产品评论数按时间分布
##总评词频、差评词频、好评词频
##各产品好中差评分布

import pandas as pd
import matplotlib.pyplot as plt
import re
from wordcloud import WordCloud
from PIL import Image
import numpy as np
import docx

top5_country = [] #买家最多的前五个国家
logistics_text = [] #最常用使用的前五个物流方式
star_count = []
rate = []
url1 = []
attr_name = []
attr_box = []
stopwords = open('static/stopwords.txt').readlines()
stopwords_list = []
name = input('文件名：')
for i in stopwords:
    clean_i = re.sub('\\n','',i)
    stopwords_list.append(clean_i)

def read_csv(filename=name+'_data.xls'):
    csv_file = pd.read_excel(filename).drop_duplicates()
    return csv_file

def logistics():
    logistics_data = read_csv()['logistics']
    Series_logistics = pd.Series(logistics_data.value_counts())
    labels = Series_logistics.index[0:5]
    values = Series_logistics.values[0:5]
    others = int(sum(Series_logistics.values)-sum(values))
    labels1 = []
    values1 = []
    for i in labels:
        labels1.append(i)
    labels1.append('others')
    for i in values:
        values1.append(int(i))
    values1.append(others)
    for i in labels:
        logistics_text.append(i)
    print(values1,labels1)
    plt.figure()
    plt.title('Proportion of logistics')
    plt.pie(values1,labels=labels1,autopct='%1.1f%%')
    plt.axis('equal')
    plt.savefig('Proportion of logistics.jpg')


def country_to_logistics():
    country = read_csv()['user_country']
    logistics = read_csv()['logistics']
    cl = pd.Series(logistics.values,index=country.values)
    fir_cl = cl[cl.index==top5_country[0]].value_counts().index[0]
    sec_cl = cl[cl.index==top5_country[1]].value_counts().index[0]
    thr_cl = cl[cl.index==top5_country[2]].value_counts().index[0]
    four_cl = cl[cl.index==top5_country[3]].value_counts().index[0]
    fif_cl = cl[cl.index==top5_country[4]].value_counts().index[0]
    list = []
    list.append(fir_cl)
    list.append(sec_cl)
    list.append(thr_cl)
    list.append(four_cl)
    list.append(fif_cl)
    dic_cl = dict(zip(top5_country,list))
    return dic_cl

def review_star():
    review_star = read_csv()['user_star']
    values = []
    values.append(len(review_star[review_star == 5]))
    values.append(len(review_star[review_star == 4]))
    values.append(len(review_star[review_star == 3]))
    values.append(len(review_star[review_star == 2]))
    values.append(len(review_star[review_star == 1]))
    labels = [5,4,3,2,1]
    for i in values:
        star_count.append(i)
    GR = (values[0]+values[1])/(values[0]+values[1]+values[2]+values[3]+values[4])
    PR = (values[3]+values[4])/(values[0]+values[1]+values[2]+values[3]+values[4])
    rate.append(GR)
    rate.append(PR)
    ##星级饼图
    plt.figure()
    plt.pie(values,labels=labels,colors=plt.cm.Reds(np.linspace(0,0.5,len(labels))),autopct='%1.1f%%',shadow=True)
    plt.axis('equal')
    plt.savefig('Star_Pie.jpg')
    ##星级条形图
    plt.figure()
    a = 6
    for x,y in zip(labels,values):
        a -= 1
        plt.text(a,y+0.05,'%d' %y,ha='center',va='bottom')
    plt.bar(labels,values)
    #plt.show()
    plt.savefig('Star_Bar.jpg')

def country():
    country = read_csv()['user_country']
    Series_country = pd.Series(country.value_counts())
    labels = Series_country.index[0:10]
    values = Series_country.values[0:10]
    labels1 = []
    values1 = []
    others = sum(Series_country.values)-sum(values)
    for i in labels:
        labels1.append(i)
    labels1.append('others')
    for i in values:
        values1.append(i)
    values1.append(others)
    for i in range(5):
        top5_country.append(labels[i])
    print(labels1,values1)
    plt.figure()
    plt.pie(values1,labels=labels1,colors=plt.cm.PuBu(np.linspace(0,0.5,len(labels1))),autopct='%1.1f%%')
    plt.title('Proportion of Buyer Country')
    plt.axis('equal')
    #plt.show()
    plt.savefig('Proportion of Buyer Country.jpg')

def review_time():
    time_list = []
    ftime_list = []
    c_time = read_csv()['user_review_time']
    ts_df = pd.DataFrame({'time':c_time.values,'star':read_csv()['user_star'].values}).sort_values(by='time')
    c_time = ts_df['time']
    urlid = read_csv()['urlid']
    b_time = []
    a_time = []
    ##处理日期
    for i in ts_df['time'].values:
        b_time.append(i[0:7])
    for i in range(len(c_time)):
        try:
            if c_time[i][0:7] not in a_time:
                a_time.append(c_time[i][0:7])
            clean_time = re.sub('\-',',',c_time[i]) #根据原数据时间格式不同，需要调整这里  ##若使用二期文件需要把 \- 改为 \?\?
            clean_time = clean_time[0:-8]           #同上
            clean_time = re.sub(' ','',clean_time)  ##如果使用二期导出的文件需要注销这两行
            time_list.append(clean_time)
        except:
            print('Format error')
    time_star = pd.Series(b_time,index=ts_df['star'].values)
    aver = []
    a_time.sort()
    summ = []
    lenn = []
    for i in a_time:
        aver1 = sum(time_star[time_star.values == i].index)/len(time_star[time_star.values == i])
        summ.append(sum(time_star[time_star.values == i].index))
        lenn.append(len(time_star[time_star.values == i]))
        aver.append(aver1)
    print(b_time)
    month_rev = []
    print(time_star)
    for i in a_time:
        month_rev.append(len(time_star[time_star.values == i]))
    print('==========================================================================================')
    print(a_time)
    print('==========================================================================================')
    print(aver)
    print('==========================================================================================')
    print(month_rev)
    print('==========================================================================================')
    print(summ)
    print('==========================================================================================')
    print(lenn)
    print('==========================================================================================')

    fig = plt.figure(figsize=(7,3))
    ax1 = fig.add_subplot(111)
    ax1.plot(a_time[1:],month_rev[1:],'b')
    ax1.set_ylabel('Total of Review')
    ax1.legend(['Total of review'],loc=2)
    ax1.set_ylim(0,max(month_rev)+1000)
    for i in range(len(month_rev)):
        if i < len(month_rev)-1:
            ax1.text(i,month_rev[i+1]-70,month_rev[i+1])
    ax2 = ax1.twinx()
    ax2.plot(a_time[1:],aver[1:],'r')
    ax2.set_ylabel('Average of Star')
    ax2.legend(['Average of star'],loc=1)
    ax2.set_ylim(min(aver)-0.02,max(aver)+0.02)
    for i in range(len(month_rev)):
        if i < len(month_rev)-1:
            ax2.text(i,aver[i+1]+0.007,'%.2f' %(aver[i+1]))
    plt.savefig('review_Line.jpg')
    '''
    for i in range(len(time_list)):
        year = int(time_list[i].split(',')[0])
        month = int(time_list[i].split(',')[1])
        day = int(time_list[i].split(',')[2])
        time = datetime.date(year,month,day)
        #time = time.strftime('%Y-%m-%d')
        ftime_list.append(time) ##ftime_list 是经过格式化处理的时间
    ##处理产品id
    goods_list = []
    np_urlid = np.array(urlid.values[0:len(ftime_list)],dtype='int64')
    Series_urlid = pd.Series(np_urlid)
    goods_count = Series_urlid.value_counts()
    for i in range(len(goods_count.index)):
        goods = goods_count.values[i]
        goods_list.append(goods)  ##goods_list是各产品的评论数量
    #print(pd.Series(ftime_list[0:goods_list[0]]))
    ##统计时间段各产品评论数
    review_sum = pd.Series(ftime_list[0:goods_list[0]]).value_counts()
    review_sum = pd.Series(review_sum).sort_index()
    #print(review_sum)
    review_sum = pd.Series(ftime_list[goods_list[0]:goods_list[0]+goods_list[1]]).value_counts()
    review_sum = pd.Series(review_sum).sort_index()
    #print(goods_list)
    x_list = []
    y_list = []
    for i in range(len(goods_list)-1):
        rev = pd.Series(ftime_list[goods_list[i]:(goods_list[i]+goods_list[i+1])]).value_counts()
        rev = pd.Series(rev).sort_index()
        x_list.append(rev.index)
        y_list.append(rev.values)
    ##绘图
    months = mdates.MonthLocator()
    days = mdates.DayLocator()
    timeFmt = mdates.DateFormatter('%Y-%m')
    fig,ax = plt.subplots()
    all_time = pd.Series(ftime_list).value_counts()
    all_time = all_time.sort_index()
    #plt.figure(figsize=(15,8))
    #print(all_time.index)
    #print(all_time.values)
    
    plt.plot(all_time.index,all_time.values,color='black')
    ax.xaxis.set_major_locator(months)
    ax.xaxis.set_major_formatter(timeFmt)
    ax.xaxis.set_minor_locator(days)
    plt.legend(['all review'],loc=2)
    plt.grid()
    #plt.show()
    plt.savefig('review_Line.jpg')
    '''


'''
TODO:调整词云大小、颜色形状等
TODO:最终输出总评词频、好评词频、差评词频
DOWNDATE:01.03 
DOWN!!!
'''

def wordcloud():
    pic_np = np.array(Image.open('static/yun.jpeg'))
    all_review_list = []
    good_review_list = []
    bad_review_list = []
    review = pd.Series(read_csv()['user_review'].values,index=read_csv()['user_star'].values).dropna().sort_index()

    good_review_5 = review[review.index == 5]
    good_review_4 = review[review.index == 4]
    bad_review_1 = review[review.index == 1]
    bad_review_2 = review[review.index == 2]

    for i in range(len(review)):
        all_review_list.append(review.values[i])
    for i in range(len(bad_review_1)):
        bad_review_list.append(bad_review_1.values[i])
    for i in range(len(bad_review_2)):
        bad_review_list.append(bad_review_2.values[i])


    for i in range(len(good_review_5)):
        good_review_list.append(good_review_5.values[i])
    for i in range(len(good_review_4)):
        good_review_list.append(good_review_4.values[i])
    with open('all_review.txt','w') as f:
        f.write(''.join(all_review_list))
    with open('good_review.txt','w') as f:
        f.write(''.join(good_review_list))
    with open('bad_review.txt','w') as f:
        f.write(''.join(bad_review_list))

    text = open('all_review.txt').read()    ##总评
    text2 = open('good_review.txt').read()  ##好评
    text3 = open('bad_review.txt').read()   ##差评

    wordcould = WordCloud(background_color='white',stopwords=stopwords_list,width=1000,height=500,mask=pic_np).generate(text)
    wordcould2 = WordCloud(background_color='white',stopwords=stopwords_list,width=1000,height=500,mask=pic_np).generate(text2)
    wordcould3 = WordCloud(background_color='white',stopwords=stopwords_list,width=1000,height=500,mask=pic_np).generate(text3)

    plt.figure(dpi=250)
    plt.title('all_review_wordcloud')
    plt.imshow(wordcould,interpolation='bilinear')
    plt.axis('off')
    plt.savefig('all_review.jpg')

    plt.figure(dpi=250)
    plt.title('good_review_wordcloud')
    plt.imshow(wordcould2,interpolation='bilinear')
    plt.axis('off')
    plt.savefig('good_review.jpg')

    plt.figure(dpi=250)
    plt.title('bad_review.wordcloud')
    plt.imshow(wordcould3,interpolation='bilinear')
    plt.axis('off')
    plt.savefig('bad_review.jpg')


def country_to_star():
    country = read_csv()['user_country']
    Series_country = pd.Series(country.value_counts())
    country_list = Series_country.index[0:5]
    cs = pd.DataFrame(read_csv()['user_country'].values,index=read_csv()['user_star'].values)
    five_star = []
    four_star = []
    three_star = []
    two_star = []
    one_star = []
    five_cs = cs[cs.index == 5]
    four_cs = cs[cs.index == 4]
    three_cs = cs[cs.index == 3]
    two_cs = cs[cs.index == 2]
    one_cs = cs[cs.index == 1]
    for i in range(len(country_list)):
        five_star.append(len(five_cs[five_cs == country_list[i]].dropna()))
        four_star.append(len(four_cs[four_cs == country_list[i]].dropna()))
        three_star.append(len(three_cs[three_cs == country_list[i]].dropna()))
        two_star.append(len(two_cs[two_cs == country_list[i]].dropna()))
        one_star.append(len(one_cs[one_cs == country_list[i]].dropna()))
    data = {'5 star':five_star,'4 star':four_star,'3 star':three_star,'2 star':two_star,'1 star':one_star}
    cs_frame = pd.DataFrame(data,index=country_list)
    return cs_frame,country_list

'''
TODO:对比加上了中评之后的国家差评率，找出差异最大的，这样可以得出哪个国家不轻易给差评，哪个国家容易给差评

'''

def count_country_to_star_rate():
    cs_frame ,country_list= country_to_star()
    rate_list_123= []
    rate_list_12 = []
    for i in range(len(country_list)):
        rate_123 = (cs_frame['1 star'][i]+cs_frame['2 star'][i]+cs_frame['3 star'][i])/sum(cs_frame.ix[i])
        rate_list_123.append(rate_123)
        rate_12 = (cs_frame['1 star'][i]+cs_frame['2 star'][i])/sum(cs_frame.ix[i])
        rate_list_12.append(rate_12)
    bad_review_rate_123 = pd.Series(rate_list_123,index=country_list).sort_values()
    country_poor_rate = dict(zip(bad_review_rate_123.index,bad_review_rate_123.values))

    bad_review_rate_12 = pd.Series(rate_list_12,index=country_list)

    plt.figure()
    plt.title('bad_review_rate(1 star+2 star+3 star)')
    plt.bar(bad_review_rate_123.index,bad_review_rate_123.values)
    plt.savefig('bad_review_rate(1 star+2 star+3 star).jpg')
    plt.figure()
    plt.title('bad_review_rate(1 star+2 star) of country')
    plt.bar(bad_review_rate_12.index,bad_review_rate_12.values)
    plt.savefig('bad_review_rate(1 star+2 star).jpg')
    return country_poor_rate

def url_poor_rate():
    urlid = read_csv()['urlid']
    star = read_csv()['user_star']
    data = pd.Series(star.values,index=urlid.values)
    url_list = []
    poor_rate_list = []
    good_rate_list = []
    for i in urlid.value_counts().index:
        url_list.append(i)
    for i in range(len(url_list)):
        try:
            url = data[url_list[i]]
            url_PR = (len(url[url==1])+len(url[url==2]))/len(url)
            url_GR = (len(url[url==5])+len(url[url==4]))/len(url)
        except:
            print(url_list[i])
        poor_rate_list.append(url_PR)
        good_rate_list.append(url_GR)
    PR = pd.Series(poor_rate_list,index=url_list).sort_values(ascending=False)
    GR = pd.Series(good_rate_list,index=url_list).sort_values(ascending=False)
    for i in PR.index:
        url1.append(i)
    for i in GR.index:
        url1.append(i)
    pr_dic = dict(zip(PR.index,PR.values))
    gr_dic = dict(zip(GR.index,PR.values))


    urlr = []
    for i in PR.index:
        urlr.append(str(i))
    plt.figure(figsize=(7,7))
    plt.title('Poor_rate')
    plt.bar(urlr[0:5],PR.values[0:5])
    plt.savefig('Poor_Rate.jpg')
    return pr_dic,gr_dic


'''
TODO:用replace()方法把产品id转换为可读性更强的id
TODO:打包成类
'''
def get_attr():
    attr = read_csv()['item_attr']
    attr1 = []
    attr2 = []
    attr3 = []
    attr_name = []

    for i in attr:
        i = re.sub(r',',':',i)
        length = len(i.split(':'))
        i = i.split(':')
        if length == 2:
            attr1.append(i[1])
            if re.sub(r'[\'\"\[]+','',i[0]) not in attr_name:
                attr_name.append(re.sub(r'[\'\"\[]+','',i[0]))
        elif length == 4:
            attr1.append(i[1])
            if re.sub(r'[\'\"\[]+','',i[0]) not in attr_name:
                attr_name.append(re.sub(r'[\'\"\[]+','',i[0]))
            attr2.append(i[3])
            if re.sub(r'[\'\"\[\] ]+','',i[2]) not in attr_name:
                attr_name.append(re.sub(r'[\'\"\]\[ ]+','',i[2]))
        elif length == 6:
            attr1.append(i[1])
            if re.sub(r'[\'\"\[]+','',i[0]) not in attr_name:
                attr_name.append(re.sub(r'[\'\"\[]+','',i[0]))
            attr2.append(i[3])
            if re.sub(r'[\'\"\[\] ]+','',i[2]) not in attr_name:
                attr_name.append(re.sub(r'[\'\"\]\[ ]+','',i[2]))
            attr3.append(i[5])
            if re.sub(r'[\'\"\[]+', '', i[4]) not in attr_name:
                attr_name.append(re.sub(r'[\'\"\[]+', '', i[4]))

    attr11 = []
    attr22 = []
    attr33 = []
    if  len(attr_name) == 3 :
        for i in attr1:
            i = re.sub(r'[\'\"\]]+', '', i)
            i = i.lower()
            attr11.append(i)
        for i in attr2:
            i = re.sub(r'[\'\"\]]+', '', i)
            i = i.lower()
            attr22.append(i)
        for i in attr33:
            i = re.sub(r'[\'\"\]]+', '', i)
            i = i.lower()
            attr33.append(i)
        S_attr1 = pd.Series(attr11).value_counts()
        S_attr2 = pd.Series(attr22).value_counts()
        plt.figure()
        labels = S_attr1.index[0:5]
        values = S_attr1.values[0:5]
        others = sum(S_attr1.values) - sum(values)
        labels1 = []
        values1 = []
        prop1_values = []
        for i in labels:
            labels1.append(i)
        labels1.append('others')
        for i in values:
            values1.append(i)
        values1.append(others)
        for i in values1:
            prop1_values.append(i / sum(S_attr1.values))
        plt.pie(values1, labels=labels1, autopct='%1.1f%%')
        plt.title('Proportion of '+attr_name[0])
        plt.axis('equal')
        # plt.show()
        plt.savefig('Proportion of '+attr_name[0] + '.jpg')

        plt.figure()
        labels = S_attr2.index
        values = S_attr2.values
        others = sum(S_attr2.values) - sum(values)
        labels2 = []
        values2 = []
        prop2_values = []
        for i in labels:
            labels2.append(i)
        labels2.append('others')
        for i in values:
            values2.append(i)
        values2.append(others)
        for i in values2:
            try:
                prop2_values.append(i / sum(S_attr2.values))
            except:
                print(S_attr2)
                print(attr_name)
        plt.pie(values2, labels=labels2, autopct='%1.1f%%')
        plt.title('Proportion of '+attr_name[1])
        plt.axis('equal')
        # plt.show()
        plt.savefig('Proportion of '+attr_name[1] + '.jpg')
        return labels1, labels2, attr_name, prop1_values, prop2_values
    elif len(attr_name)== 2 and attr_name[1] != 'ShipsFrom':
        for i in attr1:
            i = re.sub(r'[\'\"\]]+', '', i)
            i = i.lower()
            attr11.append(i)
        for i in attr2:
            i = re.sub(r'[\'\"\]]+', '', i)
            i = i.lower()
            attr22.append(i)
        for i in attr33:
            i = re.sub(r'[\'\"\]]+', '', i)
            i = i.lower()
            attr33.append(i)
        S_attr1 = pd.Series(attr11).value_counts()
        S_attr2 = pd.Series(attr22).value_counts()
        plt.figure()
        labels = S_attr1.index[0:5]
        values = S_attr1.values[0:5]
        others = sum(S_attr1.values) - sum(values)
        labels1 = []
        values1 = []
        prop1_values = []
        for i in labels:
            labels1.append(i)
        labels1.append('others')
        for i in values:
            values1.append(i)
        values1.append(others)
        for i in values1:
            prop1_values.append(i / sum(S_attr1.values))
        plt.pie(values1, labels=labels1, autopct='%1.1f%%')
        plt.title('Proportion of '+attr_name[0])
        plt.axis('equal')
        # plt.show()
        plt.savefig('Proportion of '+attr_name[0] + '.jpg')
        plt.figure()
        labels = S_attr2.index
        values = S_attr2.values
        others = sum(S_attr2.values) - sum(values)
        labels2 = []
        values2 = []
        prop2_values = []
        for i in labels:
            labels2.append(i)
        labels2.append('others')
        for i in values:
            values2.append(i)
        values2.append(others)
        for i in values2:
            try:
                prop2_values.append(i / sum(S_attr2.values))
            except:
                print(attr_name)
        plt.pie(values2, labels=labels2, autopct='%1.1f%%')
        plt.title('Proportion of '+attr_name[1])
        plt.axis('equal')
        # plt.show()
        plt.savefig('Proportion of '+attr_name[1] + '.jpg')
        return labels1, labels2, attr_name, prop1_values, prop2_values

    else:
        for i in attr1:
            i = re.sub(r'[\'\"\]]+', '', i)
            i = i.lower()
            attr11.append(i)
        S_attr1 = pd.Series(attr11).value_counts()
        plt.figure()
        labels = S_attr1.index[0:5]
        values = S_attr1.values[0:5]
        others = sum(S_attr1.values) - sum(values)
        labels1 = []
        values1 = []
        prop1_values = []
        for i in labels:
            labels1.append(i)
        labels1.append('others')
        for i in values:
            values1.append(i)
        values1.append(others)
        for i in values1:
            prop1_values.append(i / sum(S_attr1.values))
        plt.pie(values1, labels=labels1, autopct='%1.1f%%')
        plt.title('Proportion of '+attr_name[0])
        plt.axis('equal')
        # plt.show()
        plt.savefig('Proportion of '+attr_name[0] + '.jpg')
        return labels1,attr_name,prop1_values







    '''new_attr = []
    for i in range(len(attr)):
        clean_attr = re.sub(r',',':',attr[i])
        if len(clean_attr.split(':')) == 4:
            new_attr.append(clean_attr)
    attr_count = len(new_attr[0].split(':'))/2
    if attr_count == 2:
        attr1_name = re.search(r'\w+',new_attr[0].split(':')[0]).group(0)
        attr2_name = re.search(r'\w+',new_attr[0].split(':')[2]).group(0)
        attr_name.append(attr1_name)
        attr_name.append(attr2_name)
        attr1 = []
        attr2 = []
        for i in range(len(new_attr)):
            attr1.append(re.match(r'\w+',new_attr[i].split(':')[1]).group(0))
            attr2.append(re.match(r'\w+',new_attr[i].split(':')[3]).group(0))
        S_attr1 = pd.Series(attr1).value_counts()
        S_attr2 = pd.Series(attr2).value_counts()
        plt.figure()
        labels = S_attr1.index[0:5]
        values = S_attr1.values[0:5]
        others = sum(S_attr1.values)-sum(values)
        labels1 =[]
        values1 =[]
        prop1_values = []
        for i in labels:
            labels1.append(i)
        labels1.append('others')
        for i in values:
            values1.append(i)
        values1.append(others)
        for i in values1:
            prop1_values.append(i/sum(S_attr1.values))
        plt.pie(values1,labels=labels1,autopct='%1.1f%%')
        plt.title('Proportion of '+attr1_name)
        plt.axis('equal')
        #plt.show()
        plt.savefig('Proportion of '+attr1_name+'.jpg')
        plt.figure()
        labels = S_attr2.index[0:5]
        values = S_attr2.values[0:5]
        others = sum(S_attr2.values)-sum(values)
        labels2 = []
        values2 = []
        prop2_values = []
        for i in labels:
            labels2.append(i)
        labels2.append('others')
        for i in values:
            values2.append(i)
        values2.append(others)
        for i in values2:
            prop2_values.append(i/sum(S_attr2.values))
        plt.pie(values2,labels=labels2,autopct='%1.1f%%')
        plt.title('Proportion of '+attr2_name)
        plt.axis('equal')
        #plt.show()
        plt.savefig('Proportion of '+attr2_name+'.jpg')
        print(attr_count,
              '\nlabels1:',labels1,
              '\nvalues1:',values1,
              '\nattr_name:',attr_name,
              '\nprop1_values:',prop1_values,
              '\nprop2_values:',prop2_values)
        '''




def get_docx():
    dic_cl = country_to_logistics()
    country_poor_rate = count_country_to_star_rate()
    pr_dic,gr_dic = url_poor_rate()
    try:
        labels1,labels2,attr_name,prop1_values,prop2_values = get_attr()
    except:
        labels1,attr_name,prop1_values = get_attr()
    logistics = '''从上面的物流占比图可以看出，卖家使用的物流方式主要集中在%s 、%s、%s、%s 和 %s，各物流方式的占比如上图所示''' % (logistics_text[0],logistics_text[1],logistics_text[2],logistics_text[3],logistics_text[4])

    if dic_cl[top5_country[1]] == dic_cl[top5_country[2]] == dic_cl[top5_country[0]]:
        logistics_country = '''而在本次样本中，%s物流成为最多卖家的首选，对%s、%s和%s等国家都不约而同地采用上述物流方式，这很大可能是受产品性质的影响。''' \
                        % (dic_cl[top5_country[0]], top5_country[0], top5_country[1], top5_country[2],)
    elif dic_cl[top5_country[0]] == dic_cl[top5_country[1]]:
        logistics_country = '''另外，针对不同国家，卖家会使用不同的物流方式，例如%s和%s，卖家更偏向于使用%s，而%s，则是%s。'''\
                            %(top5_country[0],top5_country[1],dic_cl[top5_country[0]],top5_country[2],dic_cl[top5_country[2]])
    elif dic_cl[top5_country[0]] == dic_cl[top5_country[2]]:
        logistics_country = '''另外，针对不同国家，卖家会使用不同的物流方式，例如%s和%s，卖家更偏向于使用%s，而%s，则是%s。'''\
                            %(top5_country[0],top5_country[2],dic_cl[top5_country[0]],top5_country[1],dic_cl[top5_country[2]])
    elif dic_cl[top5_country[1]] == dic_cl[top5_country[2]]:
        logistics_country = '''另外，针对不同国家，卖家会使用不同的物流方式，例如%s和%s，卖家更偏向于使用%s，而%s，则是%s'''\
                            %(top5_country[1],top5_country[2],dic_cl[top5_country[1]],top5_country[0],dic_cl[top5_country[0]])
    else:
        logistics_country = '''另外，针对不同国家，卖家会使用不同的物流方式，例如对于 %s ，卖家使用最多的是 %s ;而 %s，%s的使用率最高；对于 %s，则是 %s , 各位卖家可以以此为参考，优化自己的物流系统。''' \
                        % (top5_country[0], dic_cl[top5_country[0]], top5_country[1], dic_cl[top5_country[1]], top5_country[2],
           dic_cl[top5_country[2]])
    country = '''从买家国家占比图可以看到，排名前五的国家分别是 %s 、%s 、 %s 、%s 和 %s ，各位卖家可以针对这三个国家进行差异化营销，设计相应语言的海报、具有国家特色的文案等，增强品牌好感；另外，从差评率的角度来看，买家最多的前五个国家中，%s 差评率为 %.2f%% ，%s 差评率为%.2f%%，%s 差评率为%.2f%%，%s 差评率为%.2f%%和%s的差评率为%.2f%%。''' \
              % (top5_country[0], top5_country[1], top5_country[2], top5_country[3], top5_country[4], top5_country[0],
           country_poor_rate[top5_country[0]]*100, top5_country[1], country_poor_rate[top5_country[1]]*100, top5_country[2],
           country_poor_rate[top5_country[2]]*100, top5_country[3], country_poor_rate[top5_country[3]]*100, top5_country[4],
           country_poor_rate[top5_country[4]]*100)
    star = '''从上图可以看到，本次分析样本的五星好评有%s条，四星好评有%s条，三星中评有%s条，二星差评有%s条，一星差评有%s条。总体好评率为%.2f%%（五星+四星），差评率为%.2f%%（二星+一星），其中，产品id为%s的样本差评率最高，为%.2f%%，其次是%s，差评率为%.2f%%。样本产品词云和差评词云如下图所示：''' \
           %(star_count[0],star_count[1],star_count[2],star_count[3],star_count[4],rate[0]*100,rate[1]*100,url1[0],pr_dic[url1[0]]*100,url1[1],pr_dic[url1[1]]*100)
    if len(attr_name) == 2 and attr_name[1] != 'ShipsFrom':
        attr = '''买家在购买商品时选择的属性是卖家或厂家备货时重要参考因素，如上图所示，在所分析的行业样本中，有%s和%s属性可供选择，其中%s属性中有%.2f%%的消费者选择了%s，而有%.2f%%的消费者选择了%s和%.2f%%的消费者选择了%s；在%s属性中，有%.2f%%的消费者选择了%s,有%.2f%%的消费者选择了%s和%.2f%%的消费者选择了%s。'''\
               % (attr_name[0],attr_name[1],attr_name[0],prop1_values[0]*100,labels1[0],prop1_values[1]*100,labels1[1],prop1_values[2]*100,labels1[2],attr_name[1],prop2_values[0]*100,labels2[0],prop2_values[1]*100,labels2[1],prop2_values[2]*100,labels2[2])
    elif len(attr_name) == 3 and attr_name[2] != 'ShipsFrom':
        attr = '''买家在购买商品时选择的属性是卖家或厂家备货时重要参考因素，如上图所示，在所分析的行业样本中，有%s和%s属性可供选择，其中%s属性中有%.2f%%的消费者选择了%s，而有%.2f%%的消费者选择了%s和%.2f%%的消费者选择了%s；在%s属性中，有%.2f%%的消费者选择了%s,有%.2f%%的消费者选择了%s和%.2f%%的消费者选择了%s。''' \
               % (attr_name[0], attr_name[1], attr_name[0], prop1_values[0] * 100, labels1[0], prop1_values[1] * 100,
                  labels1[1], prop1_values[2] * 100, labels1[2], attr_name[1], prop2_values[0] * 100, labels2[0],
                  prop2_values[1] * 100, labels2[1], prop2_values[2] * 100, labels2[2])
    elif len(attr_name) == 1 or attr_name[1] == 'ShipsFrom':
        attr = '''买家在购买商品时选择的属性是卖家或厂家备货时重要参考因素，如上图所示，在所分析的行业样本中，只有%s属性可供选择，其中有%.2f%%的人选了%s，有%.2f%%的人选了%s，有%.2f%%的人选了%s。'''\
               % (attr_name[0],prop1_values[0]*100,labels1[0],prop1_values[1]*100,labels1[1],prop1_values[2]*100,labels1[2])
    summarize = '''总的来说，根据速卖通样本数据来看，样本好评率为%.2f%%，差评率为%.2f%%；在买家来源方面，样本中大部分买家来自%s、%s、%s、%s和%s；最后在物流方面，%s是被卖家使用得最多的物流方式，针对不同国家，卖家会使用不同的物流方式，例如对于%s,卖家使用最多的是%s;而 %s，%s的使用率最高。'''\
                %(rate[0]*100,rate[1]*100,top5_country[0], top5_country[1], top5_country[2], top5_country[3], top5_country[4],logistics_text[0],top5_country[0], dic_cl[top5_country[0]], top5_country[1], dic_cl[top5_country[1]])

    ps = '请自行补充...'

    doc = docx.Document()
    doc.add_paragraph(u'一、评论星级分析', style='Heading 1')
    doc.add_picture('Star_Bar.jpg')
    doc.add_picture('Poor_Rate.jpg')
    doc.add_paragraph(star, style=None)
    doc.add_picture('all_review.jpg')
    doc.add_picture('bad_review.jpg')
    doc.add_paragraph(u'二、买家来源分析', style='Heading 1')
    doc.add_picture('Proportion of Buyer Country.jpg')
    doc.add_picture('bad_review_rate(1 star+2 star+3 star).jpg')
    doc.add_paragraph(country,style=None)

    doc.add_paragraph(u'三、物流分析', style='Heading 1')
    doc.add_picture('Proportion of logistics.jpg')
    doc.add_paragraph(logistics,style=None)
    doc.add_paragraph(logistics_country,style=None)
    try:
        if len(attr_name) == 3:
            doc.add_paragraph(u'四、商品属性分析', style='Heading 1')
            doc.add_picture('Proportion of ' + attr_name[0] + '.jpg')
            doc.add_picture('Proportion of ' + attr_name[1] + '.jpg')
            doc.add_paragraph(attr, style=None)
        if len(attr_name) == 2 and attr_name[1] != 'ShipsFrom':
            print(attr_name)
            doc.add_paragraph(u'四、商品属性分析', style='Heading 1')
            doc.add_picture('Proportion of '+ attr_name[0]+'.jpg')
            doc.add_picture('Proportion of '+ attr_name[1]+'.jpg')
            doc.add_paragraph(attr,style=None)
        elif len(attr_name) == 1 or attr_name[1] != 'ShipsFrom':
            print(attr_name)
            doc.add_paragraph(u'四、商品属性分析', style='Heading 1')
            doc.add_picture('Proportion of '+attr_name[0]+'.jpg')
            doc.add_paragraph(attr,style=None)
        else:
            print(attr_name)
            doc.add_paragraph(u'四、商品属性分析 (系统暂时无法分析，请自行分析)', style='Heading 1')
    except:
        print(attr_name)
        doc.add_paragraph(u'四、商品属性分析 (系统暂时无法分析，请自行分析)', style='Heading 1')

    doc.add_paragraph(u'五、销售趋势及平均星级变化', style='Heading 1')
    doc.add_picture('review_line.jpg')
    doc.add_paragraph(ps,style=None)
    doc.add_paragraph(u'六、评价数据总结', style='Heading 1')
    doc.add_paragraph(summarize, style=None)


    doc.save(name+'_commentssss.docx')

'''
TODO:属性分析（get_attr）需要再写一个应对单属性的代码 DOWN
TODO：属性对应国家
!!!!!!发现重大错误，所有饼图计算错误，需要重写代码：物流、国家、星级。。。。！！！！！ 1月18日 DOWN:1.20已修改
TODO:评价数据深度挖掘
'''

'''
TODO:国家与评论星级之间的关系
DOWN IN 01.04
'''
'''
TODO:根据图表数据设计模板并自动生成报告文字
DOWNDATE:01.04
'''


print('20%')
logistics()
print('25%')
review_star()
print('30%')
country()
print('35%')
country_to_logistics()
print('45%')
review_time()
print('65%')
wordcloud()
print('85%')
count_country_to_star_rate()
print('95%')
url_poor_rate()
get_attr()
get_docx()
print("100%")

print('Down!')

